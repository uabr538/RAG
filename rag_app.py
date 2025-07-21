import os
import streamlit as st
from tempfile import NamedTemporaryFile
import fitz  # PyMuPDF
import docx
import pandas as pd
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA

st.set_page_config(page_title="RAG App (Cloud Friendly)", layout="centered")
st.title("📄 Ask Your Files (RAG App - Streamlit Cloud)")

# Check for API key
try:
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    os.environ["OPENAI_API_KEY"] = openai_api_key
except KeyError:
    st.error("OpenAI API key not found. Please configure it in Streamlit Cloud secrets.")
    st.stop()

# File uploader
uploaded_files = st.file_uploader("Upload PDFs, DOCX, or Excel files", type=["pdf", "docx", "xlsx"], accept_multiple_files=True)

def load_pdf(path):
    try:
        text = ""
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        st.warning(f"Error reading PDF: {e}")
        return None

def load_docx(path):
    try:
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs if p.text])
    except Exception as e:
        st.warning(f"Error reading DOCX: {e}")
        return None

def load_excel(path):
    try:
        df = pd.read_excel(path)
        return df.to_string()
    except Exception as e:
        st.warning(f"Error reading Excel: {e}")
        return None

if uploaded_files:
    docs = []
    for f in uploaded_files:
        suffix = f.name.split(".")[-1].lower()
        with NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
            tmp.write(f.read())
            tmp_path = tmp.name

        try:
            if suffix == "pdf":
                content = load_pdf(tmp_path)
            elif suffix == "docx":
                content = load_docx(tmp_path)
            elif suffix == "xlsx":
                content = load_excel(tmp_path)
            else:
                st.warning(f"Unsupported file type: {suffix}")
                continue

            if content:
                docs.append(Document(page_content=content, metadata={"source": f.name}))
        finally:
            try:
                os.unlink(tmp_path)  # Clean up temporary file
            except Exception as e:
                st.warning(f"Error deleting temporary file {tmp_path}: {e}")

    if docs:
        question = st.text_input("Ask a question about your documents:")
        if question:
            with st.spinner("Processing documents and generating answer..."):
                try:
                    # Initialize text splitter
                    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                    split_docs = splitter.split_documents(docs)

                    # Initialize embedding model and vector store
                    embedding_model = OpenAIEmbeddings(model="text-embedding-ada-002")
                    db = Chroma.from_documents(split_docs, embedding_model, persist_directory=None)
                    retriever = db.as_retriever(search_kwargs={"k": 4})

                    # Initialize QA chain
                    qa_chain = RetrievalQA.from_chain_type(
                        llm=ChatOpenAI(temperature=0, model_name="gpt-3.5-turbo"),
                        retriever=retriever,
                        chain_type="stuff"
                    )

                    # Run query
                    answer = qa_chain.run(question)
                    st.markdown("**Answer:**")
                    st.write(answer)

                    # Clean up Chroma to free memory
                    db.delete_collection()
                except Exception as e:
                    st.error(f"Error processing question: {e}")
    else:
        st.warning("No valid documents were processed. Please check your files.")
else:
    st.info("Please upload files to start.")
